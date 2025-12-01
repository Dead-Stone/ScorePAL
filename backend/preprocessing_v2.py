#preprocessing_v2.py
import json
import os
import logging
from pathlib import Path
import shutil
import zipfile
import base64
from typing import Dict, Any, List, Optional, Tuple
from PIL import Image
import io
import nbformat
import re
import tempfile
# Removing problematic imports
# from unstructured.partition.pdf import partition_pdf
# from unstructured.documents.elements import NarrativeText, Image as UnstructuredImage
import google.generativeai as genai
from prompts.answer_key_prompt import get_answer_key_prompt
from prompts.image_prompt import get_image_description_prompt
from docx import Document
import docx2txt
import uuid
import hashlib
import time
from datetime import datetime

# PDF extraction library
import fitz  # PyMuPDF

# Import Neo4j connector
try:
    from utils.neo4j_connector import Neo4jConnector
    from utils.directory_utils import ensure_directory_structure, cleanup_temp_directories
    neo4j_available = True
except ImportError:
    neo4j_available = False

# Set Tesseract path
from dotenv import load_dotenv
load_dotenv()

# Set Tesseract paths with fallbacks
# Configure Tesseract OCR path
os.environ["TESSDATA_PREFIX"] = os.getenv("TESSDATA_PREFIX", "/usr/share/tesseract-ocr/4.00/tessdata")
tesseract_path = os.getenv("TESSDATA_PREFIX", "/usr/share/tesseract-ocr/4.00/tessdata")
os.environ["PATH"] = f"{tesseract_path}:{os.environ['PATH']}"

# os.environ["TESSDATA_PREFIX"] = "C:\Program Files\Tesseract-OCR\tessdata"
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

os.environ["EXTRACT_IMAGE_BLOCK_CROP_HORIZONTAL_PAD"] = "20"
os.environ["EXTRACT_IMAGE_BLOCK_CROP_VERTICAL_PAD"] = "20"

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz) with better formatting preservation.
    This is a standalone function that can be imported by other modules.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF with preserved formatting
    """
    logger.info(f"Extracting text from PDF: {pdf_path}")
    
    try:
        import fitz  # PyMuPDF
        text = ""
        doc = fitz.open(pdf_path)
        
        # Process each page
        for page_num, page in enumerate(doc):
            # Try different extraction methods for better formatting
            page_text = ""
            
            # Method 1: Extract text blocks (preserves structure better)
            try:
                blocks = page.get_text("blocks")
                for block in blocks:
                    if len(block) >= 5 and block[4].strip():  # block[4] contains the text
                        # Clean and format the block text
                        block_text = block[4].strip()
                        # Replace multiple spaces with single space but preserve line breaks
                        block_text = re.sub(r' +', ' ', block_text)
                        page_text += block_text + "\n\n"
            except Exception as e:
                logger.warning(f"Block extraction failed on page {page_num+1}: {e}")
                
            # Method 2: If block extraction failed or gave poor results, try dict extraction
            if not page_text.strip() or len(page_text.strip()) < 50:
                try:
                    text_dict = page.get_text("dict")
                    page_text = ""
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                line_text = ""
                                for span in line.get("spans", []):
                                    if "text" in span:
                                        line_text += span["text"]
                                if line_text.strip():
                                    page_text += line_text.strip() + "\n"
                            page_text += "\n"  # Add paragraph break
                except Exception as e:
                    logger.warning(f"Dict extraction failed on page {page_num+1}: {e}")
            
            # Method 3: Fallback to simple text extraction with better formatting
            if not page_text.strip():
                try:
                    simple_text = page.get_text()
                    if simple_text.strip():
                        # Clean up the simple text extraction
                        lines = simple_text.split('\n')
                        formatted_lines = []
                        for line in lines:
                            line = line.strip()
                            if line:
                                # Try to reconstruct words that were split
                                line = re.sub(r'(\w)\s+(\w)', r'\1\2', line)
                                # But preserve intentional spaces between words
                                line = re.sub(r'([a-z])([A-Z])', r'\1 \2', line)
                                formatted_lines.append(line)
                        page_text = '\n'.join(formatted_lines)
                except Exception as e:
                    logger.warning(f"Simple extraction failed on page {page_num+1}: {e}")
            
            # Add page text to overall text
            if page_text.strip():
                text += f"Page {page_num+1}:\n{page_text}\n\n"
        
        doc.close()  # Properly close the document
        
        # Post-process the entire text for better readability
        if text.strip():
            text = _clean_extracted_text(text)
            logger.info(f"Successfully extracted text with PyMuPDF: {len(text)} characters")
            return text
        else:
            logger.warning("No text content found in PDF")
            return "No text content found in this PDF"
            
    except ImportError:
        logger.error("PyMuPDF (fitz) not available. Please install with: pip install pymupdf")
        return "Error: PyMuPDF not installed"
    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {e}")
        return f"Error extracting text from PDF: {str(e)}"


def _clean_extracted_text(text: str) -> str:
    """
    Clean and format extracted PDF text for better readability.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned and formatted text
    """
    # Fix common OCR/extraction issues
    
    # 1. Fix specific number patterns first (most specific to least specific)
    text = re.sub(r'\b0 1 4 9 5 2 4 1 8\b', r'014952418', text)
    text = re.sub(r'\b1 4 8\b', r'148', text)
    text = re.sub(r'\b2 0 0 1 : 0 d b 8 : 8 5 a 3\b', r'2001:0db8:85a3', text)
    text = re.sub(r'\bF e 8 0\b', r'Fe80', text)
    text = re.sub(r'\b1 9 2 \. 1 6 8 \. 0 \. 0\b', r'192.168.0.0', text)
    
    # 2. Fix common patterns like "1 4 8" -> "148" (general number patterns)
    text = re.sub(r'\b(\d)\s+(\d)\s+(\d)\b', r'\1\2\3', text)
    text = re.sub(r'\b(\d)\s+(\d)\b', r'\1\2', text)
    
    # 3. Fix specific word patterns (most common first)
    specific_fixes = {
        r'\bM a t e e n\b': r'Mateen',
        r'\bA h s a n\b': r'Ahsan',
        r'\bC M P E\b': r'CMPE',
        r'\bH w (\d)\b': r'Hw\1',
        r'\bH w\b': r'Hw',
        r'\bs e c\b': r'sec',
        r'\bI P v (\d)\b': r'IPv\1',
        r'\bI P v(\d)\b': r'IPv\1',
        r'\bI P\b': r'IP',
        r'\bS u b n e t t i n g\b': r'Subnetting',
        r'\bA d d r e s s\b': r'Address',
        r'\bA l l o c a t i o n\b': r'Allocation',
        r'\bQ u e s t i o n\b': r'Question',
        r'\bD e p a r t m e n t s\b': r'Departments',
        r'\bC o m p r e s s i o n\b': r'Compression',
        r'\bH i e r a r c h i c a l\b': r'Hierarchical',
        r'\bC o n s e r v a t i o n\b': r'Conservation',
        r'\bD e s i g n\b': r'Design',
        r'\bB e t t e r\b': r'Better',
        r'\bZ e l l e\b': r'Zelle',
        r'\ba s s i g n e d\b': r'assigned',
        r'\ba d d r e s s\b': r'address',
        r'\bb l o c k\b': r'block',
        r'\bp o i n t s\b': r'points',
        r'\bh a s\b': r'has',
        r'\bb e e n\b': r'been',
        r'\bt h e\b': r'the',
        r'\ba n d\b': r'and',
    }
    
    for pattern, replacement in specific_fixes.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # 4. Fix general patterns like "C M P E" -> "CMPE" (after specific fixes)
    text = re.sub(r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b', r'\1\2\3\4', text)
    text = re.sub(r'\b([A-Z])\s+([A-Z])\s+([A-Z])\b', r'\1\2\3', text)
    text = re.sub(r'\b([A-Z])\s+([A-Z])\b', r'\1\2', text)
    
    # 5. Fix common word patterns that got split (general approach)
    # Fix 3-letter combinations that are likely one word
    text = re.sub(r'\b([a-z])\s+([a-z])\s+([a-z])\b', r'\1\2\3', text)
    
    # 6. Fix mathematical expressions
    text = re.sub(r'\b2\s*\^\s*(\d+)\b', r'2^\1', text)
    text = re.sub(r'\b(\d+)\s*\+\s*(\d+)\b', r'\1+\2', text)
    text = re.sub(r'\b(\d+)\s*-\s*(\d+)\b', r'\1-\2', text)
    text = re.sub(r'\b(\d+)\s*=\s*(\d+)\b', r'\1=\2', text)
    
    # 7. Fix IPv6 and IP address patterns
    text = re.sub(r'(\w+):\s*(\w+):\s*(\w+):\s*:', r'\1:\2:\3::', text)
    text = re.sub(r'/\s*(\d+)', r'/\1', text)
    
    # 8. Clean up excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Remove triple+ line breaks
    text = re.sub(r' +', ' ', text)  # Multiple spaces to single space
    text = re.sub(r'\s*\n\s*', '\n', text)  # Clean line breaks
    
    # 9. Fix punctuation spacing
    text = re.sub(r'\s+([.,:;!?])', r'\1', text)  # Remove space before punctuation
    text = re.sub(r'([.,:;!?])\s*([a-zA-Z0-9])', r'\1 \2', text)  # Add space after punctuation
    
    # 10. Fix common sentence patterns
    text = re.sub(r'\.\s*([a-z])', lambda m: '. ' + m.group(1).upper(), text)  # Capitalize after period
    
    return text.strip()

class FilePreprocessor:
    def __init__(self, temp_dir: str = "temp_uploads", output_dir: str = "processed_uploads", 
                 custom_output_path: str = None, save_ocr_files: bool = True):
        self.directories = ensure_directory_structure()
        self.temp_dir = self.directories["temp_uploads"]
        self.extracted_dir = self.temp_dir / "extracted"
        
        # Allow custom output directory if specified
        if custom_output_path:
            self.output_dir = Path(custom_output_path)
            os.makedirs(self.output_dir, exist_ok=True)
            logger.info(f"Using custom output directory: {self.output_dir}")
        else:
            self.output_dir = self.directories["processed_uploads"]
            
        self.logger = logger
        self.save_ocr_files = save_ocr_files
        
        # Initialize Neo4j connector if available
        self.db = None
        if neo4j_available:
            self.db = Neo4jConnector()
        
        self._setup_directories()
        
    def _generate_assignment_id(self, assignment_name: str, question_text: str) -> str:
        """Generate a unique ID for an assignment based on its content."""
        content_hash = hashlib.md5(f"{assignment_name}:{question_text}:{time.time()}".encode()).hexdigest()
        return f"assignment_{content_hash[:8]}_{int(time.time())}"

    def process_files(self, 
                     submissions_zip: str, 
                     question_paper: str, 
                     answer_key_PATH: str = None, 
                     rubric: Dict[str, Any] = None,
                     assignment_name: str = None,
                     custom_output_path: str = None,
                     save_intermediate_files: bool = True) -> Dict[str, Any]:
        """
        Process submission files, question paper, and optionally an answer key.
        
        Args:
            submissions_zip: Path to the ZIP file containing submissions
            question_paper: Path to the question paper file
            answer_key_PATH: Optional path to the answer key file
            rubric: Optional rubric for grading
            assignment_name: Optional name for the assignment
            custom_output_path: Optional custom output directory path
            save_intermediate_files: Whether to save intermediate OCR files
            
        Returns:
            Dictionary with processed submissions, question text, and answer key
        """
        # Set output path if provided
        if custom_output_path:
            self.output_dir = Path(custom_output_path)
            os.makedirs(self.output_dir, exist_ok=True)
            logger.info(f"Using custom output directory for this job: {self.output_dir}")
            
        # Set whether to save OCR files
        self.save_ocr_files = save_intermediate_files
        
        try:
            # Generate assignment name if not provided
            if not assignment_name:
                assignment_name = f"Assignment_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Process the question paper
            logger.info(f"Processing question paper: {question_paper}")
            question_text = self._process_pdf(question_paper)
            
            # Process the answer key or generate one
            answer_key_text = None
            if answer_key_PATH:
                logger.info(f"Processing answer key: {answer_key_PATH}")
                answer_key_text = self._process_pdf(answer_key_PATH)
            else:
                logger.info("Generating answer key using Gemini 1.5 Flash")
                answer_key_text = self._generate_answer_key(question_text, rubric)
            
            # Process the submissions
            zip_path = Path(submissions_zip)
            logger.info(f"Extracting submissions from {zip_path}")
            submissions = self._extract_submissions(zip_path)
            logger.info(f"Extracted {len(submissions)} submissions")
            
            # Store assignment data in Neo4j if available
            assignment_id = self._generate_assignment_id(assignment_name, question_text)
            if self.db and self.db.is_connected():
                db_assignment_id = self.db.store_assignment_data(
                    assignment_name=assignment_name,
                    question_text=question_text,
                    rubric=rubric
                )
                if db_assignment_id:
                    assignment_id = db_assignment_id
                    logger.info(f"Assignment stored in Neo4j with ID: {assignment_id}")
            
            # Save processed data to files
            self._save_processed_data(
                assignment_id=assignment_id,
                assignment_name=assignment_name,
                question_text=question_text,
                answer_key=answer_key_text,
                rubric=rubric,
                submissions=submissions
            )
            
            logger.info("Files processed successfully")
            return {
                "assignment_id": assignment_id,
                "assignment_name": assignment_name,
                "submissions": submissions,
                "question": question_text,
                "answer_key": answer_key_text,
                "rubric": rubric,
                "output_dir": str(self.output_dir)
            }
        except Exception as e:
            logger.error(f"Error processing files: {str(e)}")
            raise
        finally:
            logger.info("Cleaning up temporary files")
            # Keep files in development mode, clean up in production
            if not os.getenv("DEBUG", "False").lower() == "true":
                cleanup_temp_directories(self.directories)

    def _save_processed_data(self, 
                           assignment_id: str,
                           assignment_name: str,
                           question_text: str,
                           answer_key: str,
                           rubric: Dict[str, Any],
                           submissions: Dict[str, str]) -> None:
        """
        Save processed data to files for later use.
        
        Args:
            assignment_id: Unique ID for the assignment
            assignment_name: Name of the assignment
            question_text: Text of the question
            answer_key: Text of the answer key
            rubric: Rubric for grading
            submissions: Dictionary mapping student names to submission texts
        """
        # Create a directory for this assignment
        assignment_dir = self.output_dir / assignment_id
        os.makedirs(assignment_dir, exist_ok=True)
        logger.info(f"Saving processed data to: {assignment_dir}")
        
        # Save assignment metadata
        metadata = {
            "id": assignment_id,
            "name": assignment_name,
            "created_at": datetime.now().isoformat(),
            "submission_count": len(submissions),
            "has_answer_key": answer_key is not None,
            "has_rubric": rubric is not None,
            "saved_location": str(assignment_dir)
        }
        
        with open(assignment_dir / "metadata.json", 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        logger.info(f"Saved metadata to {assignment_dir}/metadata.json")
        
        # Save question text
        with open(assignment_dir / "question.txt", 'w', encoding='utf-8') as f:
            f.write(question_text)
        logger.info(f"Saved question text to {assignment_dir}/question.txt")
        
        # Save answer key if available
        if answer_key:
            with open(assignment_dir / "answer_key.txt", 'w', encoding='utf-8') as f:
                f.write(answer_key)
            logger.info(f"Saved answer key to {assignment_dir}/answer_key.txt")
        
        # Save rubric if available
        if rubric:
            with open(assignment_dir / "rubric.json", 'w', encoding='utf-8') as f:
                json.dump(rubric, f, indent=2)
            logger.info(f"Saved rubric to {assignment_dir}/rubric.json")
        
        # Save submissions
        submissions_dir = assignment_dir / "submissions"
        os.makedirs(submissions_dir, exist_ok=True)
        
        for student_name, submission_text in submissions.items():
            # Sanitize student name to create a safe filename
            safe_name = "".join(c if c.isalnum() else "_" for c in student_name)
            submission_file = submissions_dir / f"{safe_name}.txt"
            
            try:
                with open(submission_file, 'w', encoding='utf-8') as f:
                    f.write(submission_text)
                logger.info(f"Saved submission for {student_name} to {submission_file}")
            except Exception as e:
                logger.error(f"Error saving submission for {student_name}: {e}")
        
        logger.info(f"Successfully saved all processed data to {assignment_dir}")
        
        # Create a file indicating processing is complete
        with open(assignment_dir / "processing_complete.txt", 'w') as f:
            f.write(f"Processing completed at {datetime.now().isoformat()}")
        
        return assignment_dir

    def _process_pdf(self, pdf_path: str) -> str:
        """
        Process a PDF file using PyMuPDF (fitz) only for text extraction.
        Also extracts and processes images from pages with limited text content.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text from the PDF
        """
        logger.info(f"Processing PDF: {pdf_path}")
        
        # Create a filename base for any saved outputs
        pdf_basename = os.path.basename(pdf_path)
        pdf_name_no_ext = os.path.splitext(pdf_basename)[0]
        
        try:
            import fitz  # PyMuPDF
            text = ""
            doc = fitz.open(pdf_path)
            
            # Track images for later processing
            image_paths = []
            
            # Process each page
            for page_num, page in enumerate(doc):
                # Try different extraction methods for better formatting
                page_text = ""
                
                # Method 1: Extract text blocks (preserves structure better)
                try:
                    blocks = page.get_text("blocks")
                    for block in blocks:
                        if len(block) >= 5 and block[4].strip():  # block[4] contains the text
                            # Clean and format the block text
                            block_text = block[4].strip()
                            # Replace multiple spaces with single space but preserve line breaks
                            block_text = re.sub(r' +', ' ', block_text)
                            page_text += block_text + "\n\n"
                except Exception as e:
                    logger.warning(f"Block extraction failed on page {page_num+1}: {e}")
                    
                # Method 2: If block extraction failed or gave poor results, try dict extraction
                if not page_text.strip() or len(page_text.strip()) < 50:
                    try:
                        text_dict = page.get_text("dict")
                        page_text = ""
                        for block in text_dict.get("blocks", []):
                            if "lines" in block:
                                for line in block["lines"]:
                                    line_text = ""
                                    for span in line.get("spans", []):
                                        if "text" in span:
                                            line_text += span["text"]
                                    if line_text.strip():
                                        page_text += line_text.strip() + "\n"
                                page_text += "\n"  # Add paragraph break
                    except Exception as e:
                        logger.warning(f"Dict extraction failed on page {page_num+1}: {e}")
                
                # Method 3: Fallback to simple text extraction with better formatting
                if not page_text.strip():
                    try:
                        simple_text = page.get_text()
                        if simple_text.strip():
                            # Clean up the simple text extraction
                            lines = simple_text.split('\n')
                            formatted_lines = []
                            for line in lines:
                                line = line.strip()
                                if line:
                                    # Try to reconstruct words that were split
                                    line = re.sub(r'(\w)\s+(\w)', r'\1\2', line)
                                    # But preserve intentional spaces between words
                                    line = re.sub(r'([a-z])([A-Z])', r'\1 \2', line)
                                    formatted_lines.append(line)
                            page_text = '\n'.join(formatted_lines)
                    except Exception as e:
                        logger.warning(f"Simple extraction failed on page {page_num+1}: {e}")
                
                # Add page text to overall text
                if page_text.strip():
                    text += f"Page {page_num+1}:\n{page_text}\n\n"
                
                # Extract images if text content is limited
                if len(page_text.strip()) < 100:  # If page has little text, it might be image-heavy
                    logger.info(f"Page {page_num+1} has limited text, extracting images")
                    try:
                        # Extract images from the page
                        images = page.get_images(full=True)
                        for img_index, img_info in enumerate(images):
                            xref = img_info[0]  # Get the XREF of the image
                            base_img = doc.extract_image(xref)
                            image_bytes = base_img["image"]
                            
                            # Save image to a temporary file
                            img_filename = f"{pdf_path}_page{page_num+1}_img{img_index}.png"
                            with open(img_filename, "wb") as img_file:
                                img_file.write(image_bytes)
                            
                            image_paths.append(img_filename)
                    except Exception as img_e:
                        logger.warning(f"Image extraction failed on page {page_num+1}: {img_e}")
            
            doc.close()  # Properly close the document
            
            # Process extracted images with Gemini
            for img_path in image_paths:
                try:
                    image_desc = self._process_image(img_path)
                    if image_desc:
                        text += f"\n[Image Description: {image_desc}]\n"
                    # Clean up the temporary image file
                    os.remove(img_path)
                except Exception as img_e:
                    logger.warning(f"Failed to process image {img_path}: {img_e}")
            
            # Return extracted text
            if text.strip():
                # Clean the text using the improved cleaning function
                cleaned_text = _clean_extracted_text(text)
                logger.info(f"Successfully extracted text with PyMuPDF: {len(cleaned_text)} characters")
                
                # Save the extracted text if requested
                if self.save_ocr_files:
                    self._save_extracted_text(cleaned_text, f"{pdf_name_no_ext}_pymupdf.txt")
                    
                return self._clean_pdf_text(cleaned_text)
            else:
                logger.warning("No text content found in PDF")
                return "No text content found in this PDF"
            
        except ImportError:
            logger.error("PyMuPDF (fitz) not available. Please install with: pip install pymupdf")
            return "Error: PyMuPDF not installed"
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")
            return f"Error extracting text from PDF: {str(e)}"
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean extracted PDF text by removing headers, footers, and fixing formatting issues."""
        # Remove page numbers
        text = re.sub(r'\bPage\s+\d+\s+of\s+\d+\b', '', text)
        text = re.sub(r'\bPage\s+\d+\b', '', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR/PDF extraction issues
        text = re.sub(r'([a-z])-\s*([a-z])', r'\1\2', text)  # Fix hyphenated words
        
        # Remove headers/footers (common patterns)
        text = re.sub(r'(?i)(header|footer):\s*.*?\n', '', text)
        
        # Clean up paragraphs
        paragraphs = text.split('\n')
        clean_paragraphs = []
        
        for para in paragraphs:
            if len(para.strip()) > 0:
                clean_paragraphs.append(para.strip())
        
        return '\n\n'.join(clean_paragraphs)

    def _process_notebook(self, notebook_path: str) -> str:
        """Process Jupyter notebook files"""
        try:
            with open(notebook_path, 'r', encoding='utf-8') as f:
                nb = nbformat.read(f, as_version=4)
                
            # Extract text content from code and markdown cells
            content = []
            for cell in nb.cells:
                if cell.cell_type == 'code':
                    content.append(f"```python\n{cell.source}\n```")
                elif cell.cell_type == 'markdown':
                    content.append(cell.source)
                    
            return '\n\n'.join(content)
        except Exception as e:
            logger.error(f"Error processing notebook {notebook_path}: {e}")
            return ""
        
    def _process_image(self, image_path: str) -> str:
        """
        Process an image using Gemini model to get a description.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Text description of the image
        """
        try:
            with open(image_path, 'rb') as img_file:
                img_data = base64.b64encode(img_file.read()).decode()
            
            model = genai.GenerativeModel('gemini-2.0-flash')
            prompt = get_image_description_prompt()
            response = model.generate_content([prompt, img_data])
            logger.info(f"Image processed with Gemini: {len(response.text)} characters")
            return response.text.strip()
        except Exception as e:
            logger.error(f"Image processing failed with Gemini: {e}")
            
            # Fallback to basic image info
            try:
                with Image.open(image_path) as img:
                    width, height = img.size
                    format_name = img.format
                    mode = img.mode
                    return f"Image: {width}x{height} {format_name} {mode}"
            except Exception as img_e:
                logger.error(f"Failed to get basic image info: {img_e}")
                return "Image content (could not be processed)"

    def _generate_answer_key(self, question_text: str, rubric: Dict[str, Any]) -> str:
        try:
            model = genai.GenerativeModel('gemini-2.0-flash')
            prompt = get_answer_key_prompt(question_text, rubric)
            # print(f"Answer key prompt: {prompt}")
            response = model.generate_content(prompt)
            # print(f"Answer key generated: {response.text.strip()}")
            return response.text.strip()
        except Exception as e:
            logger.error(f"Answer key generation failed: {e}")
            return ""
            
    def _process_docx(self, docx_path: str) -> str:
        """Process DOCX files with text and hyperlink extraction"""
        try:
            # Method 1: Using docx2txt for comprehensive extraction
            text = docx2txt.process(docx_path)
            
            # Method 2: Alternative manual extraction with hyperlinks
            doc = Document(docx_path)
            full_text = []
            
            # Process paragraphs with hyperlinks
            for para in doc.paragraphs:
                # Check for hyperlinks in the paragraph
                if para._element.xpath('.//w:hyperlink'):
                    for hyperlink in para.hyperlinks:
                        # Extract hyperlink text and URL
                        link_text = ' '.join([run.text for run in hyperlink.runs])
                        link_url = hyperlink.address
                        full_text.append(f"{link_text} ({link_url})")
                else:
                    full_text.append(para.text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n\n'.join(full_text) if full_text else text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            return ""
            
    def _save_image_metadata(self, image_path: str, description: str) -> str:
        try:
            metadata_dir = self.output_dir / "metadata"
            metadata_dir.mkdir(parents=True, exist_ok=True)
            metadata_file = metadata_dir / f"{Path(image_path).stem}_metadata.json"
            metadata = {
                "image_path": str(image_path),
                "description": description
            }
            with open(metadata_file, 'w') as f:
                json.dump(metadata, f, indent=2)
            return str(metadata_file)
        except Exception as e:
            logger.error(f"Error saving image metadata: {e}")
            return ""

    def _setup_directories(self):
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.extracted_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info(f"Directories created: {self.temp_dir}, {self.extracted_dir}, {self.output_dir}")

    def _save_uploaded_file(self, uploaded_file, filename: str) -> Path:
        file_path = self.temp_dir / filename
        with open(file_path, 'wb') as f:
            f.write(uploaded_file.read())
        return file_path

    def _extract_submissions(self, zip_path: Path) -> Dict[str, str]:
        """Extract student submissions from a ZIP file."""
        submissions = {}
        processed_count = 0
        failed_count = 0
        
        # Create a folder for extracted submission files if needed
        ocr_extracts_dir = self.output_dir / "extracted_ocr_files"
        if self.save_ocr_files:
            os.makedirs(ocr_extracts_dir, exist_ok=True)
        
        try:
            # Create extraction directory
            extraction_path = self.extracted_dir
            if extraction_path.exists():
                shutil.rmtree(extraction_path)
            extraction_path.mkdir(parents=True, exist_ok=True)
            
            # Extract the ZIP file
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extraction_path)
            
            # Process each file
            for file_path in self._get_document_files(extraction_path):
                try:
                    # Get student name from filename
                    file_name = file_path.name
                    student_name = self._extract_student_name(file_path)
                    
                    # Extract text from file
                    file_extension = file_path.suffix.lower()
                    submission_text = ""
                    
                    if file_extension == '.pdf':
                        submission_text = self._process_pdf(str(file_path))
                        
                        # If PDF extraction failed or returned very little text, try enhanced OCR
                        if len(submission_text.strip()) < 100 or "Error extracting text from PDF" in submission_text:
                            logger.warning(f"Initial PDF extraction produced limited text for {file_name}, trying advanced OCR...")
                            try:
                                from extraction_service_v2 import extract_pdf_text
                                ocr_text = extract_pdf_text(str(file_path))
                                if len(ocr_text.strip()) > 100:
                                    logger.info(f"Advanced OCR extraction successful for {file_name}")
                                    submission_text = ocr_text
                                    
                                    # Save the OCR text separately if requested
                                    if self.save_ocr_files:
                                        ocr_file_path = ocr_extracts_dir / f"{file_path.stem}_ocr.txt"
                                        with open(ocr_file_path, 'w', encoding='utf-8') as f:
                                            f.write(ocr_text)
                                        logger.info(f"Saved OCR extract to {ocr_file_path}")
                            except Exception as ocr_e:
                                logger.error(f"Advanced OCR extraction also failed for {file_name}: {ocr_e}")
                    elif file_extension == '.txt':
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                submission_text = f.read()
                        except UnicodeDecodeError:
                            # Try with different encodings
                            try:
                                with open(file_path, 'r', encoding='latin-1') as f:
                                    submission_text = f.read()
                            except Exception as e:
                                logger.error(f"Failed to read text file with multiple encodings: {e}")
                    elif file_extension == '.docx':
                        submission_text = self._process_docx(str(file_path))
                    elif file_extension == '.ipynb':
                        submission_text = self._process_notebook(str(file_path))
                    else:
                        logger.warning(f"Unsupported file extension: {file_extension}")
                        continue
                    
                    # Add student identifier to the beginning of submission
                    student_identifier = f"Student_{processed_count+1}_{student_name}"
                    
                    # Store submission with text
                    if submission_text and len(submission_text.strip()) > 0:
                        submissions[student_identifier] = submission_text
                        
                        # Save individual submission file if requested
                        if self.save_ocr_files:
                            submission_file_path = ocr_extracts_dir / f"{student_identifier}.txt"
                            with open(submission_file_path, 'w', encoding='utf-8') as f:
                                f.write(submission_text)
                            logger.info(f"Saved processed submission to {submission_file_path}")
                            
                        logger.info(f"Successfully processed submission for: {student_identifier}")
                        processed_count += 1
                    else:
                        logger.warning(f"Empty submission for: {student_identifier}")
                        failed_count += 1
                        
                except Exception as e:
                    logger.error(f"Error processing submission {file_path}: {e}")
                    failed_count += 1
            
            logger.info(f"Processed {processed_count} submissions")
            logger.info(f"Extracted {len(submissions)} submissions")
            
            if failed_count > 0:
                logger.warning(f"Failed to process {failed_count} submissions")
                
            return submissions
        except zipfile.BadZipFile:
            logger.error(f"Invalid ZIP file: {zip_path}")
            return {}
        except Exception as e:
            logger.error(f"Error extracting submissions: {e}")
            return {}

    def _extract_student_name(self, file_path: Path) -> str:
        """Extract student name from filename."""
        # Get base filename without extension
        name_base = file_path.stem
        
        # Clean it up (remove non-alphanumeric characters)
        name_base = ''.join(c if c.isalnum() or c in ['_', '-', ' '] else '_' for c in name_base)
        
        # Format as Student_X
        student_number = len(os.listdir(self.extracted_dir))
        return f"Student_{student_number}_{name_base[:15]}"

    def _get_document_files(self, directory: Path) -> List[Path]:
        """Get all document files in a directory."""
        document_extensions = ['.pdf', '.txt', '.docx', '.ipynb']
        document_files = []
        
        for file_path in directory.glob('**/*'):
            if file_path.is_file() and file_path.suffix.lower() in document_extensions:
                document_files.append(file_path)
        
        return document_files

    def cleanup(self):
        """Clean up temporary files."""
        try:
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
            if self.extracted_dir.exists():
                shutil.rmtree(self.extracted_dir)
            logger.info("Temporary files cleaned up")
        except Exception as e:
            logger.error(f"Error cleaning up: {e}")

    def _save_extracted_text(self, text: str, filename: str) -> None:
        """Save extracted text to a file."""
        try:
            with open(self.output_dir / filename, 'w') as f:
                f.write(text)
            logger.info(f"Saved extracted text to {self.output_dir / filename}")
        except Exception as e:
            logger.error(f"Error saving extracted text: {e}")

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from a file based on its extension."""
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return ""
            
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._process_pdf(file_path)
            elif file_extension == '.txt':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    # Try with different encodings
                    try:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            return f.read()
                    except Exception as e:
                        logger.error(f"Failed to read text file with multiple encodings: {e}")
                        return ""
            elif file_extension == '.docx':
                return self._process_docx(file_path)
            elif file_extension == '.ipynb':
                return self._process_notebook(file_path)
            else:
                logger.warning(f"Unsupported file extension: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {e}")
            return ""
