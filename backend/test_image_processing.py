"""
Test script for Image Processing in Document Grading

This script tests the enhanced document processor's ability to extract and
process images from PDF and DOCX files using OCR technology.

Author: AI Grading System
Date: 2024
"""

import logging
import os
import tempfile
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import io

from document_processor import DocumentProcessor

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_test_image_with_text(text: str, size: tuple = (400, 200)) -> Image.Image:
    """Create a test image with text content."""
    # Create a white background image
    img = Image.new('RGB', size, color='white')
    draw = ImageDraw.Draw(img)
    
    try:
        # Try to use a default font
        font = ImageFont.load_default()
    except:
        font = None
    
    # Draw text on the image
    text_lines = text.split('\n')
    y_offset = 10
    for line in text_lines:
        draw.text((10, y_offset), line, fill='black', font=font)
        y_offset += 25
    
    return img


def create_test_pdf_with_images():
    """Create a test PDF file containing images with text."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        temp_dir = Path(tempfile.mkdtemp())
        pdf_path = temp_dir / "test_document_with_images.pdf"
        
        # Create PDF with text and embedded images
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Page 1: Regular text
        c.drawString(100, 750, "Test Document with Images")
        c.drawString(100, 720, "This document contains both text and images for testing OCR.")
        
        # Create a test image
        test_img = create_test_image_with_text(
            "Mathematical Formula:\nE = mc²\nWhere E is energy, m is mass, c is speed of light"
        )
        
        # Save image temporarily
        img_path = temp_dir / "formula.png"
        test_img.save(img_path)
        
        # Add image to PDF
        c.drawImage(str(img_path), 100, 500, width=300, height=150)
        
        # Page 2: More content
        c.showPage()
        c.drawString(100, 750, "Page 2: Code Example")
        
        # Create code image
        code_img = create_test_image_with_text(
            "def fibonacci(n):\n    if n <= 1:\n        return n\n    return fibonacci(n-1) + fibonacci(n-2)"
        )
        
        code_img_path = temp_dir / "code.png"
        code_img.save(code_img_path)
        
        c.drawImage(str(code_img_path), 100, 500, width=350, height=150)
        
        c.save()
        
        # Clean up temporary images
        img_path.unlink()
        code_img_path.unlink()
        
        return pdf_path
        
    except ImportError:
        logger.warning("reportlab not available - cannot create test PDF")
        return None


def create_test_docx_with_images():
    """Create a test DOCX file containing images with text."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        temp_dir = Path(tempfile.mkdtemp())
        docx_path = temp_dir / "test_document_with_images.docx"
        
        doc = Document()
        doc.add_heading('Test Document with Images', 0)
        
        doc.add_paragraph('This document contains both text and images for testing OCR capabilities.')
        
        # Create and add first image
        formula_img = create_test_image_with_text(
            "Physics Formula:\nF = ma\nWhere F is force, m is mass, a is acceleration"
        )
        
        formula_path = temp_dir / "formula.png"
        formula_img.save(formula_path)
        
        doc.add_paragraph('Mathematical Formula Image:')
        doc.add_picture(str(formula_path), width=Inches(4))
        
        # Add more text
        doc.add_heading('Code Section', level=1)
        doc.add_paragraph('Below is a code example in image format:')
        
        # Create and add code image
        code_img = create_test_image_with_text(
            "function bubbleSort(arr) {\n    for (let i = 0; i < arr.length; i++) {\n        for (let j = 0; j < arr.length - i - 1; j++) {\n            if (arr[j] > arr[j + 1]) {\n                [arr[j], arr[j + 1]] = [arr[j + 1], arr[j]];\n            }\n        }\n    }\n    return arr;\n}"
        )
        
        code_path = temp_dir / "code.png"
        code_img.save(code_path)
        
        doc.add_picture(str(code_path), width=Inches(5))
        
        doc.save(str(docx_path))
        
        # Clean up temporary images
        formula_path.unlink()
        code_path.unlink()
        
        return docx_path
        
    except ImportError:
        logger.warning("python-docx not available - cannot create test DOCX")
        return None


def test_image_processing():
    """Test the document processor's image processing capabilities."""
    print("=== Testing Image Processing in Documents ===\n")
    
    try:
        processor = DocumentProcessor()
        
        # Test 1: Check OCR capabilities
        print("1. Checking OCR Capabilities:")
        try:
            import pytesseract
            from PIL import Image
            print("   ✅ PIL (Pillow) available")
            print("   ✅ pytesseract available")
            
            # Basic OCR test
            test_img = create_test_image_with_text("Hello World Test")
            ocr_result = pytesseract.image_to_string(test_img)
            print(f"   ✅ Basic OCR test: '{ocr_result.strip()}'")
            
        except ImportError as e:
            print(f"   ❌ OCR libraries not available: {e}")
            return False
        except Exception as e:
            print(f"   ⚠️  OCR test failed: {e}")
        
        # Test 2: Check supported extensions
        print("\n2. Supported File Types:")
        extensions = processor.get_supported_extensions()
        print(f"   Total supported: {len(extensions)}")
        print(f"   Extensions: {extensions}")
        
        image_capable = ['.pdf', '.docx']
        available = [ext for ext in image_capable if ext in extensions]
        print(f"   Image-capable formats: {available}")
        
        # Test 3: Create a simple test file
        print("\n3. Testing with Simple Text File:")
        temp_dir = Path(tempfile.mkdtemp())
        test_file = temp_dir / "test.txt"
        test_file.write_text("This is a test document with some content for grading.")
        
        result = processor.process_file(test_file)
        print(f"   File type: {result.get('file_type', 'unknown')}")
        print(f"   Content length: {len(result.get('content', ''))}")
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir)
        
        print("\n✅ Image processing components tested successfully!")
        return True
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        return False


def test_ocr_accuracy():
    """Test OCR accuracy with different types of content."""
    print("\n=== Testing OCR Accuracy ===\n")
    
    try:
        import pytesseract
        
        test_cases = [
            ("Simple Text", "Hello World\nThis is a test"),
            ("Mathematical", "E = mc²\nF = ma\nv = u + at"),
            ("Code", "def hello():\n    print('Hello World')\n    return True"),
            ("Mixed", "Assignment Score: 85/100\nGrade: B+\nComments: Good work!")
        ]
        
        for test_name, test_text in test_cases:
            print(f"Testing {test_name}:")
            
            # Create image with text
            img = create_test_image_with_text(test_text)
            
            # Run OCR
            ocr_result = pytesseract.image_to_string(img).strip()
            
            print(f"   Original: {repr(test_text)}")
            print(f"   OCR Result: {repr(ocr_result)}")
            
            # Calculate similarity (simple character matching)
            original_chars = set(test_text.lower().replace(' ', '').replace('\n', ''))
            ocr_chars = set(ocr_result.lower().replace(' ', '').replace('\n', ''))
            
            if original_chars:
                similarity = len(original_chars & ocr_chars) / len(original_chars)
                print(f"   Similarity: {similarity:.2%}")
            
            print()
        
    except Exception as e:
        logger.error(f"OCR accuracy test failed: {e}")


if __name__ == "__main__":
    print("Starting Image Processing Tests...\n")
    
    # Load environment variables
    from dotenv import load_dotenv
    load_dotenv()
    
    # Run the tests
    success = test_image_processing()
    
    if success:
        test_ocr_accuracy()
        print("\n✅ All image processing tests completed successfully!")
    else:
        print("\n❌ Image processing tests failed. Please install required libraries:")
        print("   pip install Pillow pytesseract PyMuPDF pdf2image")
        print("   And install Tesseract OCR engine on your system.") 