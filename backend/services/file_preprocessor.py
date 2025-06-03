"""
File preprocessor for extracting text from various file formats.
"""

import os
import logging
import tempfile
from pathlib import Path
import re
import shutil
import traceback

# PDF processing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# DOCX processing
try:
    import docx
except ImportError:
    docx = None

# Image processing
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

# Setup logging
logger = logging.getLogger(__name__)

class FilePreprocessor:
    """Process files and extract text content."""
    
    def __init__(self, custom_output_path=None, save_ocr_files=False):
        """
        Initialize the file preprocessor.
        
        Args:
            custom_output_path: Optional custom path to save processed files
            save_ocr_files: Whether to save OCR results
        """
        self.custom_output_path = custom_output_path
        self.save_ocr_files = save_ocr_files
        
        # Create output directory if needed
        if self.custom_output_path:
            os.makedirs(self.custom_output_path, exist_ok=True)
    
    def extract_text_from_file(self, file_path):
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return ""
            
            extension = file_path.suffix.lower()
            
            # Process based on file extension
            if extension in ['.pdf']:
                return self._extract_text_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path)
            elif extension in ['.txt', '.md', '.csv']:
                return self._extract_text_from_text_file(file_path)
            elif extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
                return self._extract_text_from_image(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return f"[Unsupported file format: {extension}]"
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error extracting text: {str(e)}]"
    
    def _extract_text_from_pdf(self, file_path):
        """Extract text from PDF file."""
        if fitz is None:
            return "[PDF extraction not available: PyMuPDF not installed]"
        
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return f"[Error extracting PDF text: {str(e)}]"
    
    def _extract_text_from_docx(self, file_path):
        """Extract text from DOCX file."""
        if docx is None:
            return "[DOCX extraction not available: python-docx not installed]"
        
        try:
            text = ""
            doc = docx.Document(file_path)
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return f"[Error extracting DOCX text: {str(e)}]"
    
    def _extract_text_from_text_file(self, file_path):
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {str(e)}")
            return f"[Error extracting text: {str(e)}]"
    
    def _extract_text_from_image(self, file_path):
        """Extract text from image using OCR."""
        if Image is None or pytesseract is None:
            return "[OCR not available: PIL or pytesseract not installed]"
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            # Save OCR results if requested
            if self.save_ocr_files and self.custom_output_path:
                output_file = Path(self.custom_output_path) / f"{file_path.stem}_ocr.txt"
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {str(e)}")
            return f"[Error extracting image text: {str(e)}]"
    
    def _generate_answer_key(self, question_text, reference_text=None):
        """
        Generate a basic answer key from question text and reference material.
        
        This is a placeholder method that should be implemented with your 
        specific logic for generating answer keys.
        """
        # This is just a placeholder implementation
        if not question_text:
            return {}
        
        # Extract questions using a simple pattern
        questions = re.findall(r'\d+\.\s+(.+?)(?=\d+\.\s+|\Z)', question_text + '999. ')
        
        # Create a basic answer key structure
        answer_key = {
            "questions": [
                {"question_number": i+1, "question_text": q.strip(), "points": 10}
                for i, q in enumerate(questions) if q.strip()
            ],
            "total_points": len(questions) * 10 if questions else 0
        }
        
        return answer_key 