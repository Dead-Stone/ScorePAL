"""
Document Processor for Multi-Format File Support

This module handles extraction and processing of various document formats
including PDF, DOC, DOCX, TXT, and other text-based files for grading.

Author: AI Grading System
Date: 2024
"""

import logging
import io
import os
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import tempfile
import mimetypes
from datetime import datetime

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PyPDF2 not available - PDF processing disabled")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available - DOCX processing disabled")

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False
    logger.warning("mammoth not available - advanced DOC processing disabled")

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("PIL/pytesseract not available - OCR processing disabled")

try:
    import pdf2image
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False
    logger.warning("pdf2image not available - PDF image extraction disabled")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not available - advanced PDF processing disabled")


class DocumentProcessor:
    """Processor for various document formats including PDF, DOC, TXT."""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt': 'text',
            '.md': 'markdown', 
            '.pdf': 'pdf',
            '.doc': 'doc',
            '.docx': 'docx',
            '.py': 'code',
            '.js': 'code',
            '.java': 'code',
            '.html': 'code'
        }
        
    def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Process a single file and extract content."""
        try:
            file_ext = file_path.suffix.lower()
            file_type = self.supported_extensions.get(file_ext, 'unknown')
            
            content, metadata = self._extract_content(file_path, file_type)
            
            return {
                "name": file_path.name,
                "path": str(file_path),
                "file_type": file_type,
                "content": content,
                "metadata": metadata,
                "processed_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {"name": file_path.name, "error": str(e), "content": ""}
    
    def _extract_content(self, file_path: Path, file_type: str) -> Tuple[str, Dict]:
        """Extract content based on file type."""
        metadata = {}
        
        if file_type == 'pdf':
            return self._extract_pdf(file_path, metadata)
        elif file_type == 'docx':
            return self._extract_docx(file_path, metadata)
        else:
            return self._extract_text(file_path, metadata)
    
    def _extract_pdf(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract PDF content including text and images with OCR."""
        content_parts = []
        
        # Method 1: Try PyMuPDF for comprehensive extraction (text + images)
        if HAS_PYMUPDF:
            try:
                content, meta = self._extract_pdf_pymupdf(file_path, metadata)
                if content and not content.startswith("["):
                    logger.info(f"PyMuPDF extraction successful for {file_path.name}")
                    return content, meta
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed for {file_path.name}: {e}")
        
        # Method 2: Try PyPDF2 + image extraction
        if HAS_PDF and HAS_PDF2IMAGE and HAS_OCR:
            try:
                content, meta = self._extract_pdf_with_ocr(file_path, metadata)
                if content and not content.startswith("["):
                    return content, meta
            except Exception as e:
                logger.warning(f"PDF + OCR extraction failed: {e}")
        
        # Method 3: Fallback to basic PyPDF2
        if HAS_PDF:
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    metadata["extraction_method"] = "PyPDF2_text_only"
                    metadata["images_processed"] = False
                    return text, metadata
            except Exception as e:
                return f"[PDF error: {str(e)}]", metadata
        
        return "[PDF processing unavailable]", metadata
    
    def _extract_pdf_pymupdf(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract PDF using PyMuPDF with image OCR."""
        try:
            doc = fitz.open(file_path)
            content_parts = []
            image_count = 0
            page_count = len(doc)  # Store page count before closing
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                
                # Extract text
                page_text = page.get_text()
                if page_text.strip():
                    content_parts.append(f"\n--- Page {page_num + 1} Text ---\n")
                    content_parts.append(page_text)
                
                # Extract images and run OCR
                if HAS_OCR:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            # Get image data
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                                
                                # Convert to PIL Image and run OCR with handwriting optimization
                                pil_image = Image.open(io.BytesIO(img_data))
                                
                                # Try multiple OCR configurations for better handwriting recognition
                                try:
                                    ocr_text = pytesseract.image_to_string(pil_image)
                                    
                                    # If poor results, try handwriting optimization
                                    if len(ocr_text.strip()) < 20:
                                        custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,:;()[]{}/<>+=*-_| '
                                        ocr_handwriting = pytesseract.image_to_string(pil_image, config=custom_config)
                                        
                                        # Try dense text mode
                                        psm_config = r'--oem 3 --psm 4'
                                        ocr_dense = pytesseract.image_to_string(pil_image, config=psm_config)
                                        
                                        # Use the longest result
                                        all_results = [ocr_text, ocr_handwriting, ocr_dense]
                                        ocr_text = max(all_results, key=len)
                                except Exception as ocr_e:
                                    logger.warning(f"OCR failed: {ocr_e}")
                                    ocr_text = ""
                                
                                if ocr_text.strip():
                                    content_parts.append(f"\n--- Page {page_num + 1} Image {img_index + 1} (OCR) ---\n")
                                    content_parts.append(ocr_text)
                                    image_count += 1
                            
                            pix = None  # Free memory
                        except Exception as e:
                            logger.warning(f"Error processing image {img_index} on page {page_num}: {e}")
            
            # Set metadata before closing document
            metadata["extraction_method"] = "PyMuPDF_with_OCR"
            metadata["page_count"] = page_count
            metadata["images_processed"] = image_count
            metadata["has_images"] = image_count > 0
            
            doc.close()
            
            return "\n".join(content_parts), metadata
            
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")
            metadata["extraction_method"] = "PyMuPDF_failed"
            metadata["error"] = str(e)
            return f"[PyMuPDF extraction failed: {str(e)}]", metadata
    
    def _extract_pdf_with_ocr(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract PDF using PyPDF2 + pdf2image + OCR."""
        content_parts = []
        
        # Extract text with PyPDF2
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    content_parts.append(f"\n--- Page {page_num + 1} Text ---\n")
                    content_parts.append(page_text)
        
        # Extract images with pdf2image and OCR
        try:
            pages = pdf2image.convert_from_path(file_path)
            image_count = 0
            
            for page_num, page_image in enumerate(pages):
                # Run OCR on the entire page image with handwriting optimization
                try:
                    ocr_text = pytesseract.image_to_string(page_image)
                    
                    # If poor results, try handwriting optimization
                    if len(ocr_text.strip()) < 20:
                        custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,:;()[]{}/<>+=*-_| '
                        ocr_handwriting = pytesseract.image_to_string(page_image, config=custom_config)
                        
                        # Try dense text mode
                        psm_config = r'--oem 3 --psm 4'
                        ocr_dense = pytesseract.image_to_string(page_image, config=psm_config)
                        
                        # Use the longest result
                        all_results = [ocr_text, ocr_handwriting, ocr_dense]
                        ocr_text = max(all_results, key=len)
                except Exception as ocr_e:
                    logger.warning(f"Page {page_num} OCR failed: {ocr_e}")
                    ocr_text = ""
                
                # Only add if OCR found substantial new content
                if ocr_text.strip() and len(ocr_text.strip()) > 50:
                    content_parts.append(f"\n--- Page {page_num + 1} OCR Content ---\n")
                    content_parts.append(ocr_text)
                    image_count += 1
            
            metadata["extraction_method"] = "PyPDF2_plus_OCR"
            metadata["images_processed"] = image_count
            metadata["has_images"] = image_count > 0
            
        except Exception as e:
            logger.warning(f"OCR processing failed: {e}")
            metadata["extraction_method"] = "PyPDF2_text_only"
            metadata["images_processed"] = 0
        
        return "\n".join(content_parts), metadata
    
    def _extract_docx(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract DOCX content including text, tables, and images with OCR."""
        if not HAS_DOCX:
            return "[DOCX processing unavailable]", metadata
            
        try:
            doc = Document(file_path)
            content_parts = []
            image_count = 0
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                content_parts.append(f"\n--- Table {table_num + 1} ---")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content_parts.append(row_text)
            
            # Extract images with OCR
            if HAS_OCR:
                try:
                    image_count = self._extract_docx_images_ocr(doc, content_parts)
                except Exception as e:
                    logger.warning(f"DOCX image OCR failed: {e}")
            
            content = "\n".join(content_parts)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["images_processed"] = image_count
            metadata["has_images"] = image_count > 0
            metadata["extraction_method"] = "python-docx_with_OCR" if image_count > 0 else "python-docx"
            
            return content, metadata
            
        except Exception as e:
            return f"[DOCX error: {str(e)}]", metadata
    
    def _extract_docx_images_ocr(self, doc, content_parts: List[str]) -> int:
        """Extract and OCR images from DOCX document."""
        image_count = 0
        
        try:
            # Access the document's relationships to find images
            from docx.oxml.ns import qn
            from docx.image.image import Image as DocxImage
            import zipfile
            
            # Open the DOCX as a ZIP file to access images
            with zipfile.ZipFile(doc._part.package._package_file_path, 'r') as zip_file:
                # List all image files in the DOCX
                image_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        # Extract image data
                        img_data = zip_file.read(img_file)
                        
                        # Convert to PIL Image and run OCR with handwriting optimization
                        pil_image = Image.open(io.BytesIO(img_data))
                        
                        try:
                            ocr_text = pytesseract.image_to_string(pil_image)
                            
                            # If poor results, try handwriting optimization
                            if len(ocr_text.strip()) < 20:
                                custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,:;()[]{}/<>+=*-_| '
                                ocr_handwriting = pytesseract.image_to_string(pil_image, config=custom_config)
                                
                                # Try dense text mode
                                psm_config = r'--oem 3 --psm 4'
                                ocr_dense = pytesseract.image_to_string(pil_image, config=psm_config)
                                
                                # Use the longest result
                                all_results = [ocr_text, ocr_handwriting, ocr_dense]
                                ocr_text = max(all_results, key=len)
                        except Exception as ocr_e:
                            logger.warning(f"DOCX image OCR failed: {ocr_e}")
                            ocr_text = ""
                        
                        if ocr_text.strip():
                            content_parts.append(f"\n--- Image {image_count + 1} (OCR) ---\n")
                            content_parts.append(ocr_text)
                            image_count += 1
                            
                    except Exception as e:
                        logger.warning(f"Error processing DOCX image {img_file}: {e}")
                        
        except Exception as e:
            logger.warning(f"Error accessing DOCX images: {e}")
        
        return image_count
    
    def _extract_text(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract plain text content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(), metadata
        except Exception as e:
            return f"[Text error: {str(e)}]", metadata
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_extensions.keys())
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported."""
        return file_path.suffix.lower() in self.supported_extensions
    
    def process_multiple_files(self, file_paths: List[Path]) -> Dict[str, Any]:
        """
        Process multiple files and create a comprehensive analysis.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Dictionary containing comprehensive multi-file analysis
        """
        logger.info(f"Processing {len(file_paths)} files")
        
        analyses = []
        total_content = []
        file_types = set()
        languages = set()
        total_size = 0
        
        for file_path in file_paths:
            analysis = self.process_file(file_path)
            analyses.append(analysis)
            
            if not analysis.get('error', False):
                total_content.append(f"\n{'='*60}\nFILE: {analysis['name']}\n{'='*60}\n")
                total_content.append(analysis['content'])
                file_types.add(analysis['file_type'])
                total_size += len(analysis['content'])
                
                if 'language' in analysis:
                    languages.add(analysis['language'])
        
        # Determine overall file type for grading strategy
        overall_type = self._determine_overall_type(file_types)
        
        comprehensive_analysis = {
            "file_count": len(file_paths),
            "successful_files": len([a for a in analyses if not a.get('error', False)]),
            "failed_files": len([a for a in analyses if a.get('error', False)]),
            "total_size": total_size,
            "file_types": list(file_types),
            "languages": list(languages),
            "overall_type": overall_type,
            "comprehensive_content": "\n".join(total_content),
            "individual_analyses": analyses,
            "processed_at": datetime.now().isoformat(),
            "metadata": {
                "primary_type": overall_type,
                "has_code": "code" in file_types,
                "has_documents": any(t in file_types for t in ['pdf', 'docx', 'doc', 'text']),
                "has_data": "data" in file_types,
                "dominant_language": max(languages, key=lambda x: sum(1 for a in analyses if a.get('language') == x)) if languages else None
            }
        }
        
        logger.info(f"Processed {len(file_paths)} files - Overall type: {overall_type}")
        return comprehensive_analysis
    
    def _determine_overall_type(self, file_types: set) -> str:
        """
        Determine the overall submission type based on file types present.
        
        Args:
            file_types: Set of detected file types
            
        Returns:
            Overall type for grading strategy selection
        """
        if not file_types:
            return "unknown"
        
        # Priority order for type determination
        if "code" in file_types:
            return "code"
        elif any(t in file_types for t in ["pdf", "docx", "doc"]):
            return "document"
        elif "text" in file_types or "markdown" in file_types:
            return "text"
        elif "data" in file_types:
            return "data"
        else:
            return "mixed"
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_extensions.keys())
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported."""
        return file_path.suffix.lower() in self.supported_extensions 